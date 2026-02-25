"""
Module 1: Resume Intelligence Engine
=====================================
Handles PDF/DOCX extraction, NER with spaCy, and semantic skill-matching
via sentence-transformers cosine similarity.

Output: ResumeProfile dataclass consumed by the LLM orchestrator and GUI.
"""

from __future__ import annotations

import re
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import numpy as np

logger = logging.getLogger(__name__)

# ─── Lazy imports to avoid slow startup ───────────────────────────────────────

def _load_spacy():
    import spacy
    try:
        return spacy.load("en_core_web_sm")
    except OSError:
        logger.warning("spaCy model not found. Run: python -m spacy download en_core_web_sm")
        return None

def _load_sentence_transformer():
    from sentence_transformers import SentenceTransformer
    # ~80MB, downloaded once, cached in ~/.cache/torch/sentence_transformers
    return SentenceTransformer("all-MiniLM-L6-v2")


# ─── Known tech/skill vocabulary for augmented extraction ─────────────────────

TECH_VOCABULARY = {
    "languages": [
        "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust",
        "kotlin", "swift", "scala", "r", "matlab", "sql", "bash", "php", "ruby"
    ],
    "frameworks": [
        "react", "angular", "vue", "django", "flask", "fastapi", "spring", "express",
        "tensorflow", "pytorch", "keras", "sklearn", "scikit-learn", "pandas", "numpy",
        "huggingface", "langchain", "nextjs", "nodejs"
    ],
    "cloud": [
        "aws", "gcp", "azure", "docker", "kubernetes", "terraform", "ci/cd",
        "github actions", "jenkins", "ansible"
    ],
    "concepts": [
        "machine learning", "deep learning", "nlp", "computer vision", "reinforcement learning",
        "microservices", "rest api", "graphql", "agile", "scrum", "devops", "mlops",
        "data engineering", "data science", "system design", "distributed systems"
    ]
}

ALL_TECH_SKILLS = [s for group in TECH_VOCABULARY.values() for s in group]

# ─── Job role skill blueprints (for cosine similarity matching) ───────────────

JOB_ROLE_BLUEPRINTS = {
    "Software Engineer": [
        "algorithms", "data structures", "object-oriented programming", "system design",
        "REST API", "databases", "version control", "testing", "CI/CD"
    ],
    "ML Engineer": [
        "machine learning", "deep learning", "Python", "TensorFlow", "PyTorch",
        "model deployment", "MLOps", "data pipelines", "statistics", "feature engineering"
    ],
    "Data Scientist": [
        "statistical analysis", "Python", "SQL", "data visualization", "machine learning",
        "A/B testing", "pandas", "scikit-learn", "Jupyter", "business intelligence"
    ],
    "Frontend Engineer": [
        "React", "TypeScript", "CSS", "HTML", "state management", "webpack",
        "accessibility", "performance optimization", "testing", "responsive design"
    ],
    "Backend Engineer": [
        "REST API", "databases", "system design", "microservices", "Docker",
        "message queues", "caching", "authentication", "SQL", "NoSQL"
    ],
    "DevOps Engineer": [
        "Kubernetes", "Docker", "Terraform", "CI/CD", "AWS", "monitoring",
        "shell scripting", "networking", "security", "infrastructure as code"
    ],
    "Full Stack Engineer": [
        "React", "Node.js", "databases", "REST API", "Docker", "TypeScript",
        "authentication", "system design", "testing", "cloud deployment"
    ],
}


# ─── Data Structures ──────────────────────────────────────────────────────────

@dataclass
class SkillMatch:
    skill: str
    score: float          # 0.0 – 1.0 cosine similarity
    evidence: list[str]   # raw resume sentences that support this skill

@dataclass
class ResumeProfile:
    raw_text: str
    candidate_name: str
    contact_info: dict
    education: list[str]
    work_experience: list[dict]    # [{role, company, duration, bullets}]
    detected_skills: list[str]
    skill_matches: list[SkillMatch]
    overall_readiness: float       # 0–100
    job_role: str
    summary_for_llm: str           # Compact prompt-injection-safe summary
    years_experience: int = 0
    red_flags: list[str] = field(default_factory=list)


# ─── Extraction Layer ─────────────────────────────────────────────────────────

class TextExtractor:
    """Handles PDF and DOCX extraction with fallback strategies."""

    @staticmethod
    def extract_pdf(path: Path) -> str:
        try:
            import fitz  # PyMuPDF
            text_chunks = []
            with fitz.open(str(path)) as doc:
                for page in doc:
                    try:
                        text_chunks.append(page.get_text("text"))
                    except Exception:
                        # Some pages may fail (scanned images, etc.)
                        continue
            text = "\n".join(text_chunks)
            if text.strip():
                return text
        except Exception as e:
            logger.warning(f"PyMuPDF failed: {e}, trying fallback...")

        # Fallback: try pdfplumber (handles more PDF variants)
        try:
            import pdfplumber
            text_chunks = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_chunks.append(page_text)
            return "\n".join(text_chunks)
        except ImportError:
            raise RuntimeError(
                "PDF extraction failed. Install pdfplumber as fallback: pip install pdfplumber"
            )
        except Exception as e:
            raise RuntimeError(f"Could not extract text from PDF: {e}")

    @staticmethod
    def extract_docx(path: Path) -> str:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables (skills tables are common in resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)

    @classmethod
    def extract(cls, path: str | Path) -> str:
        path = Path(path)
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return cls.extract_pdf(path)
        elif suffix in (".docx", ".doc"):
            return cls.extract_docx(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}. Use PDF or DOCX.")


# ─── NLP Processing Layer ─────────────────────────────────────────────────────

class ResumeParser:
    """
    Parses raw resume text using spaCy NER + regex heuristics.
    Extracts structured candidate information.
    """

    EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}")
    PHONE_RE = re.compile(r"[\+]?[\d\s\-().]{10,17}")
    LINKEDIN_RE = re.compile(r"linkedin\.com/in/[\w-]+", re.I)
    GITHUB_RE = re.compile(r"github\.com/[\w-]+", re.I)
    YEAR_RE = re.compile(r"\b(19|20)\d{2}\b")
    DURATION_RE = re.compile(
        r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[\w\s,]*\d{4}"
        r"\s*[-–—to]+\s*(present|current|(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[\w\s,]*\d{4})",
        re.I
    )

    def __init__(self):
        self._nlp = None

    @property
    def nlp(self):
        if self._nlp is None:
            self._nlp = _load_spacy()
        return self._nlp

    def extract_contact(self, text: str) -> dict:
        return {
            "email": (self.EMAIL_RE.search(text) or [None])[0],
            "phone": (self.PHONE_RE.search(text) or [None])[0],
            "linkedin": (self.LINKEDIN_RE.search(text) or [None])[0],
            "github": (self.GITHUB_RE.search(text) or [None])[0],
        }

    def extract_name(self, text: str) -> str:
        """Extract candidate name from first few lines using NER + heuristic."""
        if self.nlp is None:
            return self._heuristic_name(text)
        first_block = text[:500]
        doc = self.nlp(first_block)
        persons = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
        if persons:
            return persons[0]
        return self._heuristic_name(text)

    def _heuristic_name(self, text: str) -> str:
        """First non-empty line with 2–4 title-cased words is likely the name."""
        for line in text.split("\n")[:10]:
            line = line.strip()
            words = line.split()
            if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w.isalpha()):
                return line
        return "Candidate"

    def extract_skills_vocab(self, text: str) -> list[str]:
        """Match against known tech vocabulary (fast, deterministic)."""
        text_lower = text.lower()
        found = []
        for skill in ALL_TECH_SKILLS:
            pattern = r"\b" + re.escape(skill.lower()) + r"\b"
            if re.search(pattern, text_lower):
                found.append(skill)
        return list(set(found))

    def extract_education(self, text: str) -> list[str]:
        edu_keywords = ["bachelor", "master", "phd", "b.sc", "m.sc", "b.e", "m.e",
                        "b.tech", "m.tech", "mba", "degree", "university", "college", "institute"]
        lines = text.split("\n")
        edu_lines = []
        for line in lines:
            if any(kw in line.lower() for kw in edu_keywords):
                edu_lines.append(line.strip())
        return edu_lines[:5]  # Top 5 education entries

    def estimate_years_experience(self, text: str) -> int:
        years = [int(y) for y in self.YEAR_RE.findall(text)]
        if len(years) >= 2:
            span = max(years) - min(years)
            return max(0, min(span, 30))
        return 0

    def extract_work_experience(self, text: str) -> list[dict]:
        """Lightweight work experience extractor using section detection."""
        sections = re.split(
            r"\n(?=(?:experience|work history|employment|positions?|career)\b)",
            text, flags=re.I
        )
        experience_blocks = []
        if len(sections) > 1:
            exp_text = sections[1]
            # Split on duration patterns or company/role lines
            jobs = re.split(r"\n(?=[A-Z][A-Za-z\s&,]+\n)", exp_text)
            for job in jobs[:6]:
                if len(job.strip()) > 20:
                    lines = [l.strip() for l in job.split("\n") if l.strip()]
                    experience_blocks.append({
                        "raw": job.strip()[:300],
                        "title_line": lines[0] if lines else "",
                        "bullets": [l for l in lines[1:] if l.startswith(("•", "-", "·", "*"))]
                    })
        return experience_blocks


# ─── Semantic Matching Layer ──────────────────────────────────────────────────

class SemanticMatcher:
    """
    Uses sentence-transformers to produce cosine similarity scores
    between resume text and job role skill requirements.
    """

    def __init__(self):
        self._model = None

    @property
    def model(self):
        if self._model is None:
            logger.info("Loading sentence-transformers model (first run may take ~30s)...")
            self._model = _load_sentence_transformer()
        return self._model

    def match_skills(self, resume_text: str, job_role: str) -> list[SkillMatch]:
        """
        For each skill in the job role blueprint, compute cosine similarity
        against the full resume text and extract supporting sentences.
        """
        from sklearn.metrics.pairwise import cosine_similarity

        blueprint = JOB_ROLE_BLUEPRINTS.get(job_role, JOB_ROLE_BLUEPRINTS["Software Engineer"])

        # Encode all skill descriptions + full resume in one batch call
        all_texts = blueprint + [resume_text]
        embeddings = self.model.encode(all_texts, convert_to_numpy=True, show_progress_bar=False)

        skill_embeddings = embeddings[:-1]
        resume_embedding = embeddings[-1].reshape(1, -1)

        similarities = cosine_similarity(resume_embedding, skill_embeddings)[0]

        # Find supporting sentences for each skill
        sentences = [s.strip() for s in re.split(r"[.\n]", resume_text) if len(s.strip()) > 20]

        matches = []
        for skill, score in zip(blueprint, similarities):
            # Find top supporting sentence
            if sentences:
                sent_embeddings = self.model.encode(sentences, convert_to_numpy=True, show_progress_bar=False)
                skill_emb = self.model.encode([skill], convert_to_numpy=True, show_progress_bar=False)
                sent_scores = cosine_similarity(skill_emb, sent_embeddings)[0]
                top_idx = np.argsort(sent_scores)[-2:][::-1]
                evidence = [sentences[i] for i in top_idx if sent_scores[i] > 0.3]
            else:
                evidence = []

            matches.append(SkillMatch(
                skill=skill,
                score=float(np.clip(score, 0, 1)),
                evidence=evidence
            ))

        return sorted(matches, key=lambda x: x.score, reverse=True)

    def compute_readiness(self, matches: list[SkillMatch]) -> float:
        if not matches:
            return 0.0
        scores = [m.score for m in matches]
        # Weighted: top half matters more
        top_half = scores[:len(scores)//2]
        bottom_half = scores[len(scores)//2:]
        weighted = (np.mean(top_half) * 0.7 + np.mean(bottom_half) * 0.3) if bottom_half else np.mean(top_half)
        return round(float(weighted) * 100, 1)


# ─── LLM Summary Builder ──────────────────────────────────────────────────────

def build_llm_summary(profile: "ResumeProfile") -> str:
    """
    Build a compact, injection-safe candidate summary for the LLM system prompt.
    Keeps token usage minimal while preserving maximum signal.
    """
    top_skills = [m.skill for m in profile.skill_matches[:6] if m.score > 0.4]
    weak_skills = [m.skill for m in profile.skill_matches[-4:] if m.score < 0.35]

    exp_blocks = []
    for exp in profile.work_experience[:3]:
        exp_blocks.append(f"  - {exp.get('title_line', 'Role')} | {exp.get('raw', '')[:100]}")

    summary = f"""CANDIDATE PROFILE (auto-extracted):
Name: {profile.candidate_name}
Target Role: {profile.job_role}
Estimated Experience: {profile.years_experience} years
Overall Readiness Score: {profile.overall_readiness}/100

Strong Skills ({len(top_skills)}): {', '.join(top_skills) if top_skills else 'Not detected'}
Skill Gaps ({len(weak_skills)}): {', '.join(weak_skills) if weak_skills else 'None detected'}

Education: {'; '.join(profile.education[:2]) if profile.education else 'Not specified'}

Recent Experience:
{chr(10).join(exp_blocks) if exp_blocks else '  - Not parsed'}

Detected Technologies: {', '.join(profile.detected_skills[:15]) if profile.detected_skills else 'None'}
"""
    return summary.strip()


# ─── Public API ───────────────────────────────────────────────────────────────

class ResumeIntelligenceEngine:
    """
    Main entry point for Module 1.

    Usage:
        engine = ResumeIntelligenceEngine()
        profile = engine.process("resume.pdf", job_role="ML Engineer")
    """

    def __init__(self):
        self.extractor = TextExtractor()
        self.parser = ResumeParser()
        self.matcher = SemanticMatcher()

    def process(self, file_path: str | Path, job_role: str = "Software Engineer",
                progress_callback=None) -> ResumeProfile:
        def _report(pct, msg):
            if progress_callback:
                progress_callback(pct, msg)

        logger.info(f"Processing resume: {file_path} for role: {job_role}")

        # Step 1: Extract raw text
        _report(5, "Extracting text from document...")
        raw_text = self.extractor.extract(file_path)
        if len(raw_text.strip()) < 50:
            raise ValueError("Resume appears to be empty or unreadable. Try a different file.")
        _report(15, "Text extracted successfully")

        # Step 2: Parse structured info
        _report(20, "Identifying candidate name...")
        name = self.parser.extract_name(raw_text)
        _report(25, "Extracting contact information...")
        contact = self.parser.extract_contact(raw_text)
        _report(30, "Parsing education history...")
        education = self.parser.extract_education(raw_text)
        _report(35, "Analyzing work experience...")
        work_exp = self.parser.extract_work_experience(raw_text)
        _report(42, "Detecting technical skills...")
        detected_skills = self.parser.extract_skills_vocab(raw_text)
        _report(48, "Estimating years of experience...")
        years_exp = self.parser.estimate_years_experience(raw_text)

        # Step 3: Semantic skill matching
        _report(55, "Loading AI models for skill matching...")
        logger.info("Running semantic skill matching (sentence-transformers)...")
        skill_matches = self.matcher.match_skills(raw_text, job_role)
        _report(78, "Computing readiness score...")
        readiness = self.matcher.compute_readiness(skill_matches)

        # Step 4: Detect red flags
        _report(85, "Analyzing red flags...")
        red_flags = []
        if years_exp == 0:
            red_flags.append("No work experience detected")
        if not detected_skills:
            red_flags.append("No technical skills detected — possible non-standard formatting")
        if readiness < 30:
            red_flags.append(f"Low role match ({readiness}%) — significant skill gaps")

        # Step 5: Build profile
        _report(90, "Building candidate profile...")
        profile = ResumeProfile(
            raw_text=raw_text,
            candidate_name=name,
            contact_info=contact,
            education=education,
            work_experience=work_exp,
            detected_skills=detected_skills,
            skill_matches=skill_matches,
            overall_readiness=readiness,
            job_role=job_role,
            years_experience=years_exp,
            red_flags=red_flags,
            summary_for_llm=""
        )

        # Step 6: Build LLM prompt summary (after profile is complete)
        _report(95, "Generating interview briefing...")
        profile.summary_for_llm = build_llm_summary(profile)

        _report(100, "Resume analysis complete!")
        logger.info(f"Resume processed. Readiness: {readiness}% | Skills: {len(detected_skills)}")
        return profile


# ─── Available Job Roles ──────────────────────────────────────────────────────

AVAILABLE_JOB_ROLES = list(JOB_ROLE_BLUEPRINTS.keys())
